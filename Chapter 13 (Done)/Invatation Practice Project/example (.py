# invatation.py - creates format for an invatation letter and adds names to it

import docx, os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.shared import Pt

os.chdir(r'C:\Coding Shit\myPrograms\ATBS Chapter 13\Invatation Practice Project')

namesText = open('guests.txt')
names = namesText.read()
names = names.split('\n')

doc = docx.Document()
stylesDoc = docx.Document('demo.docx')

paraObj1 = doc.add_paragraph('It would be a pleasure to have the company of')
paraObj1.style = 'SuperFancy'
