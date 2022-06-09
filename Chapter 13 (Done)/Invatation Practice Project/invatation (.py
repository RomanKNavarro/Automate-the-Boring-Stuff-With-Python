# invatation.py - creates format for an invatation letter and adds names to it

import docx, os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.shared import Pt

os.chdir(r'C:\Coding Shit\myPrograms\ATBS Chapter 13\Invatation Practice Project')

namesText = open('guests.txt')
names = namesText.read()
names = names.split('\n')

doc = docx.Document()
stylesDoc = docx.Document("C:\Coding Shit\myPrograms\ATBS Chapter 13\Invatation Practice Project\demo.docx")


for name in names:
   paraObj1 = doc.add_paragraph('It would be a pleasure to have the company of')
   
   paraObj2 = doc.add_paragraph(name)
   run = paraObj2.runs[0]
   run.bold = True
   
   paraObj3 = doc.add_paragraph('at 11010 Memory Lane on the Evening of')

   paraObj4 = doc.add_paragraph('April 1')
   run = paraObj4.runs[0]
   run.bold = True

   paraObj5 = doc.add_paragraph("at 7 o'clock'")
   paraObj5.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

for paraNum in range(0, len(doc.paragraphs)):
   paraObj = doc.paragraphs[paraNum]
   paraObj.alignment = WD_ALIGN_PARAGRAPH.CENTER
   paraObj.style = stylesDoc.styles['Super Fancy']
    

doc.save('invitations.docx')
